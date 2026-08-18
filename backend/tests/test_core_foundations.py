import os
import sys
import uuid
import tempfile
import unittest
from datetime import datetime, timezone
import fitz  # PyMuPDF
import docx
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
from sqlalchemy.pool import StaticPool
from fastapi.testclient import TestClient

sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "../..")))

from backend.database import Base, get_db
from backend.models.document import Document
from backend.models.conversation import Conversation, Message
from backend.utils.hashing import calcular_hash
from backend.utils.pdf import extraer_texto_pdf
from backend.utils.docx import extraer_texto_docx
from backend.utils.chunking import dividir_texto, crear_chunks_con_metadata
from backend.main import app

TEST_DATABASE_URL = "sqlite:///:memory:"
test_engine = create_engine(
    TEST_DATABASE_URL,
    connect_args={"check_same_thread": False},
    poolclass=StaticPool
)
TestingSessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=test_engine)


def override_get_db():
    db = TestingSessionLocal()
    try:
        yield db
    finally:
        db.close()


class TestCoreFoundations(unittest.TestCase):

    @classmethod
    def setUpClass(cls):
        Base.metadata.create_all(bind=test_engine)
        app.dependency_overrides[get_db] = override_get_db
        cls.client = TestClient(app)

    @classmethod
    def tearDownClass(cls):
        Base.metadata.drop_all(bind=test_engine)

    def setUp(self):
        self.db = TestingSessionLocal()

    def tearDown(self):
        self.db.rollback()
        self.db.close()

    def test_database_and_orm_models(self):
        """Prueba 1: Creación y persistencia de registros ORM Document, Conversation y Message."""
        unique_hash = uuid.uuid4().hex + uuid.uuid4().hex
        doc = Document(
            name="test_contract.pdf",
            original_name="Contrato_Original.pdf",
            type="pdf",
            hash=unique_hash,
            size=2048,
            status="processed"
        )
        self.db.add(doc)
        self.db.commit()
        self.db.refresh(doc)

        self.assertIsNotNone(doc.id)
        self.assertEqual(doc.name, "test_contract.pdf")
        self.assertEqual(doc.status, "processed")

        # 2. Crear conversación y mensajes
        conv = Conversation()
        self.db.add(conv)
        self.db.commit()
        self.db.refresh(conv)

        msg1 = Message(conversation_id=conv.id, role="user", content="¿Cuáles son las políticas?")
        msg2 = Message(conversation_id=conv.id, role="assistant", content="Las políticas establecen...")
        self.db.add_all([msg1, msg2])
        self.db.commit()

        self.db.refresh(conv)
        self.assertEqual(len(conv.messages), 2)
        self.assertEqual(conv.messages[0].role, "user")
        self.assertEqual(conv.messages[1].role, "assistant")

    def test_calcular_hash(self):
        """Prueba 2: Cálculo de hash SHA-256 de un archivo binario."""
        with tempfile.NamedTemporaryFile(delete=False) as tmp:
            tmp.write(b"Hello PolicyLens-AI SHA-256 Test Content")
            tmp_path = tmp.name

        try:
            hash_res = calcular_hash(tmp_path)
            self.assertEqual(len(hash_res), 64)
            self.assertEqual(hash_res, calcular_hash(tmp_path))
        finally:
            if os.path.exists(tmp_path):
                os.remove(tmp_path)

        with self.assertRaises(FileNotFoundError):
            calcular_hash("/ruta/inexistente/archivo.bin")

    def test_chunking_and_metadata(self):
        """Prueba 3: Fragmentación de texto (chunking) y asignación de metadatos."""
        texto = "Palabra " * 100
        chunks = dividir_texto(texto, tamano_chunk=300, overlap=50)

        self.assertGreater(len(chunks), 1)
        for chunk in chunks:
            self.assertLessEqual(len(chunk), 300)

        texto_corto = "Texto pequeño"
        chunks_cortos = dividir_texto(texto_corto, tamano_chunk=500)
        self.assertEqual(len(chunks_cortos), 1)
        self.assertEqual(chunks_cortos[0], texto_corto)

        meta_chunks = crear_chunks_con_metadata(
            document_id=1,
            document_name="manual_rrhh.pdf",
            page=5,
            section="Vacaciones",
            chunks=chunks
        )

        self.assertEqual(len(meta_chunks), len(chunks))
        self.assertEqual(meta_chunks[0]["document_id"], 1)
        self.assertEqual(meta_chunks[0]["document_name"], "manual_rrhh.pdf")
        self.assertEqual(meta_chunks[0]["page"], 5)
        self.assertEqual(meta_chunks[0]["section"], "Vacaciones")
        self.assertEqual(meta_chunks[0]["chunk_index"], 0)

    def test_extraer_texto_pdf(self):
        """Prueba 4: Generación y extracción de texto de un archivo PDF sintético con PyMuPDF."""
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            pdf_path = tmp.name

        try:
            doc = fitz.open()
            page1 = doc.new_page()
            page1.insert_text((50, 50), "Texto de la página 1 del PDF de prueba.")
            page2 = doc.new_page()
            page2.insert_text((50, 50), "Texto de la página 2 del PDF de prueba.")
            doc.save(pdf_path)
            doc.close()

            resultado = extraer_texto_pdf(pdf_path)
            self.assertEqual(len(resultado), 2)
            self.assertEqual(resultado[0]["page"], 1)
            self.assertIn("página 1", resultado[0]["text"])
            self.assertEqual(resultado[1]["page"], 2)
            self.assertIn("página 2", resultado[1]["text"])
        finally:
            if os.path.exists(pdf_path):
                os.remove(pdf_path)

    def test_extraer_texto_docx(self):
        """Prueba 5: Generación y extracción de texto de un archivo DOCX sintético con python-docx."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            docx_path = tmp.name

        try:
            doc = docx.Document()
            doc.add_heading("Título de Prueba", level=1)
            doc.add_paragraph("Este es el primer párrafo de contenido.")
            doc.add_paragraph("Este es el segundo párrafo de contenido.")
            doc.save(docx_path)

            resultado = extraer_texto_docx(docx_path)
            self.assertEqual(len(resultado), 1)
            self.assertEqual(resultado[0]["page"], 1)
            self.assertIn("Título de Prueba", resultado[0]["text"])
            self.assertIn("primer párrafo", resultado[0]["text"])
            self.assertIn("segundo párrafo", resultado[0]["text"])
        finally:
            if os.path.exists(docx_path):
                os.remove(docx_path)

    def test_fastapi_endpoints(self):
        """Prueba 6: Endpoints / y /health de FastAPI."""
        response_root = self.client.get("/")
        self.assertEqual(response_root.status_code, 200)
        self.assertEqual(response_root.json()["status"], "online")

        response_health = self.client.get("/health")
        self.assertEqual(response_health.status_code, 200)
        self.assertEqual(response_health.json()["status"], "ok")
        self.assertIn("timestamp", response_health.json())


if __name__ == "__main__":
    unittest.main()
